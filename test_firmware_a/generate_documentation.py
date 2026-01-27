"""Generate a very detailed project documentation Word file (English).

Includes per-module explanations for both boards, pin rationale, flows,
and operational guidance. Output goes to scripts/documentation.
"""

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def _add_table(doc, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Item'
    table.rows[0].cells[1].text = 'Pins'
    table.rows[0].cells[2].text = 'Why these pins'
    for r, (item, pins, why) in enumerate(rows, start=1):
        table.rows[r].cells[0].text = item
        table.rows[r].cells[1].text = pins
        table.rows[r].cells[2].text = why


def _add_module_section(doc, title, rows):
    doc.add_heading(title, level=2)
    for name, purpose, behavior in rows:
        doc.add_paragraph(f"{name}: {purpose}", style='List Bullet')
        doc.add_paragraph(behavior)


def create_documentation():
    """Create an in-depth Word doc for new teammates (English)."""

    doc = Document()

    # Title and subtitle
    title = doc.add_heading('IoT Dual-Board ESP32 – Full Walkthrough', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Sensor Board A + Actuator Board B with Node-RED bridge')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True

    # Table of contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Big Picture and Roles',
        '2. Repository Layout and Per-Module Guide (Board A & B)',
        '3. Hardware and Pin Map (with rationale)',
        '4. Boot Sequence and Main Loop Flow',
        '5. Data and Command Flow (boards, Node-RED, app)',
        '6. Alarm Logic, States, and Actions',
        '7. Non-Blocking Design and Timers',
        '8. Debug, Logging, and Remote Control',
        '9. Simulation Mode',
        '10. Quick How-To Operate and Test',
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

    # 1. Big Picture
    doc.add_heading('1. Big Picture and Roles', level=1)
    doc.add_paragraph(
        'Two ESP32 boards cooperate: Board A collects and interprets sensor data; Board B drives actuators. '
        'Board A is also the only northbound gateway to Node-RED / the mobile app via MQTT. '
        'The split keeps sensing and actuation responsive and isolates failures.'
    )
    doc.add_paragraph('Board A (Sensors): temperature, CO, heart rate/SpO2, ultrasonic presence, buttons; alarm logic; cloud bridge.', style='List Bullet')
    doc.add_paragraph('Board B (Actuators): LEDs (G/B/R), servo gate, LCD 16x2, passive buzzer, DFPlayer audio, buttons.', style='List Bullet')
    doc.add_paragraph('Connectivity: ESP-NOW for A↔B, MQTT for A↔Node-RED/app, UDP for local debug commands.', style='List Bullet')

    # 2. Repo layout and per-module guide
    doc.add_heading('2. Repository Layout and Per-Module Guide (Board A & B)', level=1)
    doc.add_paragraph('This section walks through every key module so newcomers can map code to behavior quickly.')

    doc.add_heading('2.1 Board A (Sensors) Modules', level=2)
    _add_module_section(doc, 'Core', [
        ('main.py', 'Entry point; orchestrates init then non-blocking loop.', 'Steps: OTA → WiFi → logging → sensors init/sim → ESP-NOW → UDP → Node-RED MQTT; then loops calling update functions.'),
        ('core.sensor_loop', 'Scheduler for sensor reads and alarm evaluation.', 'Uses elapsed() per sensor interval; calls drivers, then alarm_logic.update(); no blocking sleep.'),
        ('core.state', 'Central store for readings, alarm levels, received actuator state, control flags.', 'Avoids duplicated state; tracks stale timestamps and reboot requests.'),
        ('core.timers', 'elapsed() helpers.', 'Monotonic tick-based gating for any periodic task.'),
        ('core.wifi', 'WiFi connect.', 'Used mainly for UDP logging and MQTT; only blocking during init.'),
    ])
    _add_module_section(doc, 'Sensors', [
        ('sensors.temperature (DS18B20)', 'Temperature via OneWire.', 'Stateful conversion start + read after ~750ms; updates state.sensor_data.temperature.'),
        ('sensors.co (MQ-7 analog)', 'CO ppm with fast baseline.', 'Short startup baseline, drift follow, ppm clamp; writes state.sensor_data.co.'),
        ('sensors.heart_rate (MAX30102)', 'IR/RED reflectance → BPM + SpO2.', 'SoftI2C on 21/22; finger detect; peak/BPM buffers; updates state.sensor_data.heart_rate.*'),
        ('sensors.ultrasonic (HC-SR04)', 'Distance/presence.', 'Micro-delays per shot only; smoothing; presence feeds gate automation on B.'),
        ('sensors.buttons', 'Local buttons.', 'Reads b1/b2/b3 → state.button_state; can trigger test actions.'),
        ('sensors.accelerometer', 'Optional ADC accel.', 'Safe no-op if not wired; would fill state.sensor_data.acc.*'),
        ('sensors.simulation', 'Synthetic data.', 'Used when simulate_sensors=true; keeps flows identical without hardware.'),
    ])
    _add_module_section(doc, 'Logic', [
        ('logic.alarm_logic', 'Turns readings into warning/danger.', 'Per-sensor timers (warning/danger/recovery); aggregates worst level; emits events to B via ESP-NOW.'),
    ])
    _add_module_section(doc, 'Communication', [
        ('communication.espnow_communication', 'ESP-NOW client (A initiates).', 'Sends sensor snapshots/events to B; receives ACK/actuator state; retries; marks stale after >15s.'),
        ('communication.nodered_client', 'MQTT bridge to Node-RED/app.', 'Publishes combined state; receives commands; reconnect loop; disabled gracefully if broker down.'),
        ('communication.udp_commands', 'UDP listener (local).', 'Non-blocking socket; JSON commands; hands off to command_handler.'),
        ('communication.command_handler', 'Command router.', 'Validates, executes locally, or forwards to B based on target.'),
    ])
    _add_module_section(doc, 'Debug', [
        ('debug.debug', 'Hierarchical logging + flags.', 'Prefix + wildcard matching; remote UDP optional; silent on sink failure.'),
        ('debug.remote_log', 'UDP sender for logs.', 'Tiny, non-blocking sender used by debug.'),
    ])

    doc.add_heading('2.2 Board B (Actuators) Modules', level=2)
    _add_module_section(doc, 'Core', [
        ('main.py', 'Entry point; mirrors A flow.', 'OTA → WiFi → logging → buttons/actuators init/sim → ESP-NOW → UDP; then loop.'),
        ('core.actuator_loop', 'Scheduler for actuators.', 'Updates LED blinking, servo smoothing, LCD refresh, buzzer patterns, audio status; all timed, non-blocking.'),
        ('core.state', 'Actuator-side state.', 'LED modes, servo angle/moving, buzzer/audio flags, received sensor state + stale detection.'),
        ('core.timers', 'elapsed() helper.', 'Same timing primitive as A.'),
        ('core.wifi', 'WiFi connect.', 'For UDP logging; not critical for ESP-NOW.'),
    ])
    _add_module_section(doc, 'Actuators', [
        ('actuators.leds', 'G/B/R LEDs with modes.', 'off/on/blinking; red mirrors alarm; green baseline on.'),
        ('actuators.servo', 'Gate servo smooth motion.', '50 Hz PWM; ramps toward target; gate opens on presence from A, closes after delay.'),
        ('actuators.lcd', '16x2 I2C display.', 'Default text; alarm messages; respects user overrides; non-blocking clear/restore.'),
        ('actuators.buzzer', 'Passive buzzer patterns.', 'Warning/danger patterns; phase-based non-blocking; supports mute flag.'),
        ('actuators.audio (DFPlayer)', 'Audio playback.', 'UART1 commands for play/stop/volume; for voice/tones.'),
        ('actuators.buttons', 'Local button/SOS.', 'Reads physical button; can trigger SOS/emergency.'),
        ('actuators.simulation', 'Simulated outputs.', 'Used when simulate_actuators=true; keeps logic intact without hardware.'),
    ])
    _add_module_section(doc, 'Communication', [
        ('communication.espnow_communication', 'ESP-NOW server (B listens).', 'Receives sensor data/events; sends actuator state/ACK; stale >15s detection.'),
        ('communication.udp_commands', 'UDP listener.', 'Same JSON protocol; routed via command_handler.'),
        ('communication.command_handler', 'Router.', 'Executes local actuator commands; can forward to A if needed.'),
    ])
    _add_module_section(doc, 'Debug', [
        ('debug.debug', 'Same logging system as A.', 'Hierarchical flags; remote UDP optional.'),
    ])

    # 3. Hardware and pins
    doc.add_heading('3. Hardware and Pin Map (with rationale)', level=1)
    doc.add_paragraph('Board A (sensors):')
    _add_table(doc, [
        ('DS18B20 temperature', 'GPIO4', 'Single-wire bus; avoids I2C pins; not a boot strap pin.'),
        ('MQ-7 CO analog', 'GPIO35 (ADC1)', 'ADC1 usable with WiFi; stable ref; isolated from ADC2 WiFi conflicts.'),
        ('Buttons b1/b2/b3', 'GPIO16 / 17 / 19', 'Non-strap inputs; grouped for wiring clarity.'),
        ('Ultrasonic HC-SR04', 'TRIG GPIO5, ECHO GPIO18', 'No boot conflict; echo on input-capable pin; works with time_pulse_us.'),
        ('MAX30102 heart/SpO2', 'I2C SDA 21, SCL 22', 'Standard ESP32 I2C pins; shared bus OK.'),
        ('(Optional) Accelerometer', 'Configurable ADC pins', 'Not wired now; driver safe if absent.'),
    ])

    doc.add_paragraph('Board B (actuators):')
    _add_table(doc, [
        ('LEDs (G/B/R)', 'GPIO16 / 17 / 19', 'Direct outputs; mirrors numbering style used on A for simplicity.'),
        ('Servo SG90', 'GPIO23 (PWM)', 'Stable 50 Hz PWM; isolated from comms/I2C pins.'),
        ('LCD 1602 I2C', 'SDA 21, SCL 22', 'Shared I2C; known-good pins; 3.3V safe.'),
        ('Passive buzzer', 'GPIO25 (PWM)', 'PWM-capable; not overlapping with other buses.'),
        ('DFPlayer Mini', 'TX 27, RX 26 (UART1)', 'Keeps console UART0 free; 3.3V logic.'),
        ('User/SOS button', 'GPIO18', 'Input-friendly; not a strap pin.'),
    ])

    # 4. Boot and loop
    doc.add_heading('4. Boot Sequence and Main Loop Flow', level=1)
    doc.add_paragraph('Startup (both boards, blocking allowed):')
    doc.add_paragraph('1) OTA check; 2) WiFi connect (for logging/MQTT); 3) init remote logging; 4) init hardware (or simulation); 5) init ESP-NOW; 6) start UDP listener; 7) (A only) init MQTT bridge.', style='List Bullet')
    doc.add_paragraph('Failures log warnings but do not halt the device (loops keep running).', style='List Bullet')

    doc.add_paragraph('Main loop (non-blocking):')
    doc.add_paragraph('System control checks (reboot flag).', style='List Bullet')
    doc.add_paragraph('sensor_loop.update() (A) or actuator_loop.update() (B) — internally timed.', style='List Bullet')
    doc.add_paragraph('ESP-NOW update() — heartbeats, data, ACK, stale detection.', style='List Bullet')
    doc.add_paragraph('MQTT update() (A) — publish state, process commands.', style='List Bullet')
    doc.add_paragraph('UDP commands update() — local control/testing.', style='List Bullet')
    doc.add_paragraph('No blocking sleeps; only micro-delays inside drivers when required (DS18B20 conversion, ultrasonic pulse).', style='List Bullet')

    # 5. Data and command flow
    doc.add_heading('5. Data and Command Flow (boards, Node-RED, app)', level=1)
    doc.add_paragraph('Board-to-board (ESP-NOW):')
    doc.add_paragraph('A → B: sensor snapshot every ~2.5s (distance, temp, CO, HR/SpO2, presence, alarm level, message IDs).', style='List Bullet')
    doc.add_paragraph('B → A: actuator state heartbeat; ACKs; both sides mark stale if >15s silence.', style='List Bullet')

    doc.add_paragraph('Cloud/app path (MQTT via Node-RED, handled only by A):')
    doc.add_paragraph('Feeds (config): state/event publish to esp-a-to-nodered; commands consumed from nodered-to-esp-a; optional ack to esp-a-to-nodered-ack.', style='List Bullet')
    doc.add_paragraph('App → Node-RED → MQTT command feed → Board A. If target is B, A forwards via ESP-NOW.', style='List Bullet')
    doc.add_paragraph('A publishes combined state (its sensors + latest B actuators) on interval (default 20s).', style='List Bullet')

    doc.add_paragraph('Local/network debug path (UDP):')
    doc.add_paragraph('send_command.py can inject commands (simulate sensors, drive actuators, request status).', style='List Bullet')
    doc.add_paragraph('Non-blocking sockets; safe if no listener present.', style='List Bullet')

    doc.add_paragraph('Buttons path:')
    doc.add_paragraph('Board B button triggers local SOS/emergency and is also sent to A via ESP-NOW.', style='List Bullet')

    # 6. Alarm logic
    doc.add_heading('6. Alarm Logic, States, and Actions', level=1)
    doc.add_paragraph('Sensor thresholds (configurable): temp safe 10–35°C; CO critical >50 ppm; HR low <50 / high >120 BPM; SpO2 <90%.', style='List Bullet')
    doc.add_paragraph('Time windows: CO warning 5s / danger 30s; temp warning 10s / danger 60s; HR warning 10s / danger 60s; recovery windows per sensor.', style='List Bullet')
    doc.add_paragraph('State machine: each sensor tracks timers; global alarm = worst level among sensors.', style='List Bullet')
    doc.add_paragraph('Actions on B: red LED steady (danger) or blinking (warning); buzzer warning/danger patterns; LCD alarm text; servo gate driven by presence; green LED baseline on.', style='List Bullet')
    doc.add_paragraph('Presence automation: A sets presence_detected via ultrasonic; B opens servo to 90°; closes after configurable delay when presence is lost.', style='List Bullet')

    # 7. Non-blocking design
    doc.add_heading('7. Non-Blocking Design and Timers', level=1)
    doc.add_paragraph('elapsed() wraps all periodic work; if interval not met, functions return immediately.', style='List Bullet')
    doc.add_paragraph('Sensors each have their own cadence; alarm logic has its own cadence; actuators update incrementally.', style='List Bullet')
    doc.add_paragraph('Drivers use only short, localized waits (e.g., DS18B20 conversion, ultrasonic pulse) so the main loop stays responsive.', style='List Bullet')

    # 8. Debug and logging
    doc.add_heading('8. Debug, Logging, and Remote Control', level=1)
    doc.add_paragraph('Hierarchical logging (debug/debug.py): enable/disable by prefix or wildcard.', style='List Bullet')
    doc.add_paragraph('Remote UDP logging optional; failures are silent (non-blocking).', style='List Bullet')
    doc.add_paragraph('Noise control: set_all_logs(False) then enable specific channels for focused tests.', style='List Bullet')
    doc.add_paragraph('UDP commands: fast local control without MQTT; same JSON schema as in command_handler.', style='List Bullet')

    # 9. Simulation
    doc.add_heading('9. Simulation Mode', level=1)
    doc.add_paragraph('Enable via config.json: simulate_sensors (A) or simulate_actuators (B).', style='List Bullet')
    doc.add_paragraph('Loops swap to simulation modules; comms and logic stay identical, so flows are realistic.', style='List Bullet')

    # 10. Quick how-to
    doc.add_heading('10. Quick How-To Operate and Test', level=1)
    doc.add_paragraph('Power on A and B; watch logs (WiFi/ESP-NOW).', style='List Bullet')
    doc.add_paragraph('Check ESP-NOW link: no stale warnings after 15s; A should receive actuator state; B should receive sensor state.', style='List Bullet')
    doc.add_paragraph('MQTT/Node-RED: ensure broker reachable; A publishes to esp-a-to-nodered and listens on nodered-to-esp-a.', style='List Bullet')
    doc.add_paragraph('Use send_command.py for UDP tests (CO spike, LED toggle, servo move).', style='List Bullet')
    doc.add_paragraph('Use log_listener.py to view live logs over UDP.', style='List Bullet')

    # Save document
    output_path = os.path.expanduser('~/OneDrive/Desktop/Python Files/First Project/scripts/documentation')
    os.makedirs(output_path, exist_ok=True)
    doc_path = os.path.join(output_path, 'IoT_System_Detailed_Documentation_v2.docx')
    if os.path.exists(doc_path):
        try:
            os.remove(doc_path)
        except Exception:
            pass
    doc.save(doc_path)
    print(f"Documentation created successfully: {doc_path}")
    return doc_path

if __name__ == '__main__':
    create_documentation()
