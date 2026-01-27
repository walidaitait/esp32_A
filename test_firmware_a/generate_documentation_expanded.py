"""Generate a very comprehensive 20-30 page documentation Word file (English).

Detailed walkthrough of the entire IoT system with per-module explanations,
pin rationale, flows, logic, examples, and operational guidance.
Output: scripts/documentation/IoT_System_Complete_Guide.docx
"""

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


def _add_table(doc, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Item'
    table.rows[0].cells[1].text = 'Pins'
    table.rows[0].cells[2].text = 'Why these pins'
    for r, (item, pins, why) in enumerate(rows, start=1):
        table.rows[r].cells[0].text = item
        table.rows[r].cells[1].text = pins
        table.rows[r].cells[2].text = why


def create_documentation():
    """Create a comprehensive 20-30 page guide."""

    doc = Document()

    # Title and subtitle
    title = doc.add_heading('IoT Dual-Board ESP32 System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Complete Technical Walkthrough and Operational Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True

    subtitle2 = doc.add_paragraph('For New Team Members')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2.runs[0].font.size = Pt(12)
    subtitle2.runs[0].font.italic = True

    doc.add_paragraph()

    # Table of contents
    doc.add_heading('Table of Contents', level=1)
    toc = [
        '1. Introduction and System Overview',
        '2. Architecture and Design Philosophy',
        '3. Board A (Sensor Board) – Complete Module Guide',
        '4. Board B (Actuator Board) – Complete Module Guide',
        '5. Hardware Pins, Wiring, and Pin Selection Rationale',
        '6. Initialization Flow and Boot Sequence',
        '7. Main Loop Architecture and Non-Blocking Design',
        '8. Communication Protocols and Data Flow',
        '9. Alarm Logic System – Detailed Walkthrough',
        '10. State Management and Shared Data Structures',
        '11. Debug, Logging, and Remote Monitoring',
        '12. Simulation Mode and Testing',
        '13. Node-RED Integration and MQTT Topics',
        '14. Operational Procedures and Troubleshooting',
        '15. Code Examples and Usage Patterns',
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_paragraph()

    # 1. Introduction
    doc.add_heading('1. Introduction and System Overview', level=1)
    doc.add_paragraph(
        'This system is an IoT monitoring and control platform built on two cooperating ESP32 microcontrollers. '
        'The goal is to sense environmental conditions (temperature, air quality, heart rate, presence) and react '
        'with controlled actuators (LEDs, servo gate, display, audio feedback). The system is designed to be reliable, '
        'responsive, and scalable with minimal latency and no blocking operations in the main loops.'
    )
    doc.add_paragraph(
        'Why two boards instead of one? Separating sensing from actuation offers several advantages. First, it isolates '
        'failure domains: if a sensor fails or hangs, it does not freeze actuators. Second, it enables independent scaling: '
        'you can add many more sensors to Board A or many more actuators to Board B without redesigning everything. Third, '
        'it simplifies debugging and testing by dividing responsibility. Board A focuses on data acquisition and analysis; '
        'Board B focuses on user feedback and system state visualization.'
    )
    doc.add_paragraph(
        'The two boards communicate wirelessly via ESP-NOW, a lightweight proprietary protocol from Espressif. This is faster '
        'and simpler than WiFi for direct board-to-board messaging. Board A also connects to the internet via WiFi to act as '
        'a bridge to Node-RED and a mobile app (via MQTT). This architecture creates a three-layer system: local sensing/logic, '
        'local actuation, and cloud connectivity—all on the same pair of devices.'
    )

    # 2. Architecture and design philosophy
    doc.add_heading('2. Architecture and Design Philosophy', level=1)
    doc.add_paragraph(
        'The system was designed around two core principles: non-blocking execution and centralized state. Understanding these '
        'principles is critical to working with the codebase.'
    )
    doc.add_heading('2.1 Non-Blocking Design', level=2)
    doc.add_paragraph(
        'Every operation in the main loop must complete quickly or return immediately. There are no sleep() calls, no blocking '
        'reads, and no waiting for sensors or network. Instead, timers guard periodic work. A timer mechanism (elapsed() from '
        'core.timers) checks whether enough time has passed since the last execution of a task. If not, the task returns without '
        'doing anything. This pattern repeats for every subsystem: sensor reads, alarm logic, LED updates, servo movement, '
        'communication, and more.'
    )
    doc.add_paragraph(
        'Why is this important? Because the main loop runs at full CPU speed, the system stays responsive to interrupts, '
        'commands, and time-sensitive operations. If one sensor takes a long time (e.g., DS18B20 temperature conversion), '
        'it does not freeze LED blinking, servo motion, or the ability to receive network commands. Instead, the driver uses '
        'a state machine: it starts the conversion in one iteration and reads the result in a later iteration, without blocking.'
    )
    doc.add_heading('2.2 Centralized State', level=2)
    doc.add_paragraph(
        'All system state lives in a single shared dictionary in core.state (one for each board). Instead of each module storing '
        'its own copy of, say, temperature or LED status, every module reads and writes from this central store. This eliminates '
        'synchronization problems: if the alarm logic updates a sensor reading, the communication layer sees it immediately on the '
        'next update. If a network command changes an actuator, the main display loop sees it right away.'
    )
    doc.add_paragraph(
        'The state dictionary is organized by category: sensor_data, alarm_state, actuator_state, received_sensor_state (on B), '
        'received_actuator_state (on A), and system_control. Each category holds relevant flags, values, and metadata (like timestamps '
        'for stale detection).'
    )

    # 3. Board A (Sensor Board) – Complete Module Guide
    doc.add_heading('3. Board A (Sensor Board) – Complete Module Guide', level=1)
    doc.add_paragraph(
        'Board A is the "brain" of the system. It reads sensors, evaluates alarms, and bridges to the cloud. Let us walk through '
        'every module and understand what it does and how it works.'
    )

    doc.add_heading('3.1 Main Entry Point (main.py)', level=2)
    doc.add_paragraph(
        'main.py is the first code executed when the board powers on. Its job is to orchestrate a carefully ordered initialization, '
        'then enter the main loop. Here is the sequence:'
    )
    doc.add_paragraph('Step 1: OTA (Over-The-Air) Update Check', style='List Number')
    doc.add_paragraph(
        'The firmware checks if a new version is available. This is done first, before importing anything else, so that outdated '
        'code is not loaded. If an update is pending, the board downloads and installs it, then reboots.'
    )
    doc.add_paragraph('Step 2: WiFi Connection', style='List Number')
    doc.add_paragraph(
        'The board connects to WiFi. This is necessary for UDP logging and (optionally) MQTT. The connection is blocking at init time, '
        'but it times out quickly if the network is unavailable. Once done, the rest of the firmware proceeds whether WiFi is connected '
        'or not (local operation is not blocked by network issues).'
    )
    doc.add_paragraph('Step 3: Remote Logging Initialization', style='List Number')
    doc.add_paragraph(
        'The debug system initializes a UDP sink so logs can be sent over the network to a monitoring station. This is optional and '
        'safe to fail: if no listener is present, logs just do not go anywhere.'
    )
    doc.add_paragraph('Step 4: Hardware (Sensors) Initialization or Simulation Mode', style='List Number')
    doc.add_paragraph(
        'If simulate_sensors is false in config.json, the board initializes each sensor driver (temperature, CO, heart rate, ultrasonic). '
        'If any sensor fails to initialize (e.g., I2C device not found), a warning is logged and the system continues. If simulate_sensors '
        'is true, a simulation module is loaded instead, which generates realistic synthetic readings. This allows testing without hardware.'
    )
    doc.add_paragraph('Step 5: ESP-NOW Communication Setup', style='List Number')
    doc.add_paragraph(
        'Board A becomes a client and establishes a link to Board B (server). This wireless channel is now ready to send sensor data '
        'and receive actuator feedback.'
    )
    doc.add_paragraph('Step 6: UDP Command Listener Startup', style='List Number')
    doc.add_paragraph(
        'A UDP socket is opened on port 37022 in non-blocking mode. This allows local test scripts (send_command.py) to inject commands '
        'for debugging and testing.'
    )
    doc.add_paragraph('Step 7: Node-RED / MQTT Bridge Setup', style='List Number')
    doc.add_paragraph(
        'Board A attempts to connect to the MQTT broker (e.g., Adafruit IO). This is optional: if the broker is unavailable, the firmware '
        'continues and retries periodically.'
    )
    doc.add_paragraph(
        'Once all initialization is done, main.py enters the non-blocking main loop. It runs continuously, calling sensor_loop.update(), '
        'alarm_logic.update(), espnow_communication.update(), nodered_client.update(), and udp_commands.update() in sequence. None of these '
        'functions block; they all use internal timers to decide whether to do work or return immediately.'
    )

    doc.add_heading('3.2 Sensor Loop (core.sensor_loop)', level=2)
    doc.add_paragraph(
        'sensor_loop is the coordinator for all sensor operations. It is called from the main loop on every iteration. Internally, it uses '
        'timers to decide which sensors should be read on the current iteration.'
    )
    doc.add_paragraph(
        'For example, if temperature should be read every 1000 ms, then sensor_loop calls read_temperature() only when 1000 ms have passed '
        'since the last call. Between calls, read_temperature() returns immediately without doing anything. The same logic applies to CO, '
        'heart rate, ultrasonic, and buttons: each has its own configured interval.'
    )
    doc.add_paragraph(
        'After all sensor reads are done, sensor_loop calls alarm_logic.update() to evaluate whether any sensor is in a warning or danger '
        'state. This happens on its own interval (e.g., every 200 ms).'
    )
    doc.add_paragraph(
        'The pseudocode is roughly:'
    )
    doc.add_paragraph('if elapsed("temp", TEMP_INTERVAL): read_temperature()', style='List Bullet')
    doc.add_paragraph('if elapsed("co", CO_INTERVAL): read_co()', style='List Bullet')
    doc.add_paragraph('if elapsed("heart_rate", HEART_RATE_INTERVAL): read_heart_rate()', style='List Bullet')
    doc.add_paragraph('if elapsed("ultrasonic", ULTRASONIC_INTERVAL): read_ultrasonic()', style='List Bullet')
    doc.add_paragraph('if elapsed("buttons", BUTTON_INTERVAL): read_buttons()', style='List Bullet')
    doc.add_paragraph('if elapsed("logic", LOGIC_INTERVAL): alarm_logic.update()', style='List Bullet')

    doc.add_heading('3.3 Sensor Drivers (sensors/temperature.py, sensors/co.py, etc.)', level=2)
    doc.add_paragraph(
        'Each sensor has its own driver module. Let us look at each one to understand how sensors are read without blocking the main loop.'
    )

    doc.add_heading('Temperature Sensor (DS18B20)', level=3)
    doc.add_paragraph(
        'The DS18B20 is a digital temperature sensor connected via a one-wire bus to GPIO4. To read it, the sensor must be told to start '
        'a conversion (which takes about 750 ms), then after the conversion is done, the result is read. The naive approach is to block for '
        '750 ms. But that would freeze the system. Instead, the driver uses a state machine:'
    )
    doc.add_paragraph('First call: Start conversion, record the start time.', style='List Number')
    doc.add_paragraph('Subsequent calls (before 750 ms): Check elapsed time; if < 750 ms, return immediately.', style='List Number')
    doc.add_paragraph('After 750 ms: Read the result, update state.sensor_data["temperature"].', style='List Number')
    doc.add_paragraph(
        'This way, the temperature driver does not block the main loop. Instead, it progresses over multiple iterations until the result is ready.'
    )

    doc.add_heading('Carbon Monoxide Sensor (MQ-7)', level=3)
    doc.add_paragraph(
        'The MQ-7 is an analog CO sensor; it outputs a voltage proportional to CO concentration. Board A reads this analog value on GPIO35 '
        '(ADC1 channel). The raw ADC value is converted to millivolts, then compared to a baseline. A quick baseline is taken at startup to '
        'account for sensor drift. Afterward, the CO reading is computed as the delta from baseline, scaled to ppm (parts per million).'
    )
    doc.add_paragraph(
        'Why a baseline? The MQ-7 sensor has a fixed offset when the device powers on (it is not at 0 ppm). By measuring the first few readings '
        'and treating them as baseline, we can report the CO level relative to the startup condition, even without a calibration procedure. Over time, '
        'the baseline is slightly adjusted to follow very slow drift.'
    )
    doc.add_paragraph(
        'The reading is very fast (a few milliseconds per ADC read), so no state machine is needed; the driver simply reads on schedule and updates '
        'state.sensor_data["co"].'
    )

    doc.add_heading('Heart Rate Sensor (MAX30102)', level=3)
    doc.add_paragraph(
        'The MAX30102 is a reflectance pulse oximeter: it shines red and infrared LEDs at a fingertip and measures light reflection. From these two '
        'signals, BPM (beats per minute) and SpO2 (blood oxygen saturation) are calculated. The sensor is connected via I2C (SDA/SCL on GPIO21/22) '
        'and runs at 100 samples per second internally.'
    )
    doc.add_paragraph(
        'The driver reads raw IR and RED values in a circular buffer, detects peaks in the IR signal to estimate BPM, and computes SpO2 from the '
        'ratio of AC components. A finger-detect threshold prevents erroneous readings when no finger is present. All of this is stateful (buffers, '
        'timers, baseline calibration) and progresses incrementally over multiple updates. The sensor does not "block"; data is continuously streamed '
        'at 100 Hz and processed in the driver.'
    )

    doc.add_heading('Ultrasonic Distance Sensor (HC-SR04)', level=3)
    doc.add_paragraph(
        'The HC-SR04 measures distance via ultrasound. It requires a short trigger pulse on GPIO5, then measures the duration of a high pulse on GPIO18. '
        'The time-to-distance conversion uses the speed of sound. This driver has one small blocking section: the time_pulse_us() function (a built-in '
        'MicroPython function) waits for the echo pulse, but with a timeout to prevent infinite hangs. Once the pulse time is captured, the calculation '
        'is non-blocking. Distance readings are smoothed using a running average to reduce jitter.'
    )
    doc.add_paragraph(
        'A "presence" flag is also set based on whether the distance is below a threshold (e.g., <50 cm means someone is nearby). This flag is used '
        'by Board B to automate the servo gate.'
    )

    doc.add_heading('Buttons', level=3)
    doc.add_paragraph(
        'Buttons are simple digital inputs on GPIO16, GPIO17, and GPIO19. The button driver reads the pin states on a fast interval (e.g., 50 ms) and '
        'updates state.button_state["b1"], ["b2"], ["b3"]. These can be used to trigger test actions locally or propagate commands to Board B.'
    )

    doc.add_heading('3.4 Alarm Logic (logic.alarm_logic)', level=2)
    doc.add_paragraph(
        'The alarm logic module turns raw sensor readings into a system-wide alarm state. It evaluates each sensor reading against configured thresholds '
        'and time windows to determine if the system is in normal, warning, or danger state.'
    )
    doc.add_paragraph(
        'For each sensor (CO, temperature, heart rate), the module tracks a timer. When a reading crosses a threshold, the timer starts. If the condition '
        'persists for a "warning time" (e.g., 5 seconds for CO), the sensor state becomes "warning". If it persists for a "danger time" (e.g., 30 seconds), '
        'the state becomes "danger". If the condition clears, the timer resets and the state returns to "normal" (with a "recovery time" to avoid fluttering).'
    )
    doc.add_paragraph(
        'The global alarm level is the worst of all sensor levels: if any sensor is "danger", global alarm is "danger"; else if any is "warning", global '
        'alarm is "warning"; else "normal". This state is updated in state.alarm_state and broadcast to Board B every ~2.5 seconds via ESP-NOW.'
    )
    doc.add_paragraph(
        'On Board B, the alarm state drives actuator reactions: red LED solid for danger, blinking for warning, off for normal; buzzer patterns play for '
        'warning/danger; LCD displays alarm messages.'
    )

    doc.add_heading('3.5 Communication: ESP-NOW Client (communication.espnow_communication)', level=2)
    doc.add_paragraph(
        'Board A sends sensor data and receives actuator feedback via ESP-NOW. The protocol is stateful and non-blocking. Here is the flow:'
    )
    doc.add_paragraph(
        'Every ~2.5 seconds, Board A packages the latest sensor readings (temperature, CO, heart rate, ultrasonic distance, presence, alarm state) into '
        'a message and sends it to Board B. The send operation is non-blocking; the hardware handles it asynchronously.'
    )
    doc.add_paragraph(
        'Board B receives the message, updates its own copy of the sensor state, sends an ACK (acknowledgment) back to A, and continues processing. '
        'Board B also periodically sends its own actuator state (LED modes, servo angle, etc.) to Board A.'
    )
    doc.add_paragraph(
        'Both boards track ACKs. If a board does not receive an ACK from the other for >15 seconds, the received state is marked "stale", indicating '
        'a possible link failure. This triggers a warning log so operators know communication is broken.'
    )
    doc.add_paragraph(
        'The send/receive loop is all non-blocking, so ESP-NOW communication does not interrupt sensor reading or actuator updates.'
    )

    doc.add_heading('3.6 Communication: Node-RED Bridge (communication.nodered_client)', level=2)
    doc.add_paragraph(
        'Board A connects to an MQTT broker (e.g., Adafruit IO) to act as a bridge to Node-RED and a mobile app. Every ~3-20 seconds (configurable), '
        'Board A publishes the combined system state (its own sensors + last known actuator state from B) to a feed (MQTT topic). This data appears on '
        'Node-RED dashboards and can be viewed on a smartphone app.'
    )
    doc.add_paragraph(
        'Node-RED sends commands to Board A on a separate feed. These commands are JSON-formatted and can target Board A directly or Board B (in which case '
        'Board A forwards them via ESP-NOW). Commands can be anything: toggle LEDs, move the servo, change sensor simulation values, etc. Because MQTT is '
        'built on TCP, it is reliable; messages are not lost due to network hiccups.'
    )
    doc.add_paragraph(
        'The MQTT connection is optional: if the broker is down or unreachable, Board A logs a warning and retries every ~5 seconds. The rest of the system '
        '(local sensing, esp-now communication, actuation) continues normally without MQTT.'
    )

    doc.add_heading('3.7 Communication: UDP Commands (communication.udp_commands)', level=2)
    doc.add_paragraph(
        'For debugging and testing, Board A listens for UDP commands on port 37022. A test script (send_command.py) can inject JSON commands over the local '
        'network. Examples: {"target":"A", "command":"simulate", "args":["co", "100"]} to fake a CO spike. The UDP socket is non-blocking, so incoming '
        'commands are only checked once per loop cycle.'
    )

    doc.add_heading('3.8 Command Handler (communication.command_handler)', level=2)
    doc.add_paragraph(
        'Commands from UDP, MQTT, or buttons all go through command_handler.py. It parses the command, validates it, and either executes it locally or '
        'forwards it to Board B. For example, a command to toggle the green LED is forwarded to B (since B controls the LEDs); a command to read sensor '
        'state is handled locally.'
    )

    doc.add_heading('3.9 State (core.state)', level=2)
    doc.add_paragraph(
        'core/state.py is a single Python module that holds all of Board A\'s state as module-level variables (dictionaries). Every other module imports '
        'state and reads/writes its fields. The state includes:'
    )
    doc.add_paragraph('sensor_data: current temperature, CO, heart rate, SpO2, ultrasonic distance, presence flag', style='List Bullet')
    doc.add_paragraph('system_state: per-sensor alarm levels (CO, temp, heart rate)', style='List Bullet')
    doc.add_paragraph('alarm_state: overall alarm level and source', style='List Bullet')
    doc.add_paragraph('button_state: b1, b2, b3 press flags', style='List Bullet')
    doc.add_paragraph('received_actuator_state: last known LED, servo, LCD, buzzer state from Board B', style='List Bullet')
    doc.add_paragraph('system_control: reboot flags, control commands', style='List Bullet')
    doc.add_paragraph(
        'By centralizing state, there is a single source of truth. No module has to worry about sync bugs or consistency.'
    )

    doc.add_heading('3.10 Timers (core.timers)', level=2)
    doc.add_paragraph(
        'The elapsed() function is a simple utility that tracks monotonic time and returns True only if N milliseconds have passed since the last True return '
        'for a given key. For example: if elapsed("temp", 1000): read_temperature(). First time called, it returns True. Until 1000 ms later, it returns False. '
        'This pattern is used everywhere in the codebase to enforce cadences without blocking sleep calls.'
    )

    doc.add_heading('3.11 WiFi (core.wifi)', level=2)
    doc.add_paragraph(
        'Handles WiFi connection at startup. On Board A, WiFi is needed for UDP logging and MQTT. It is not critical for the main operation (ESP-NOW works '
        'without WiFi), so if WiFi fails, the system continues. WiFi is not reconnected in the main loop; it is only set up at init time.'
    )

    doc.add_heading('3.12 Debug Logging (debug.debug)', level=2)
    doc.add_paragraph(
        'All log messages from any module go through debug.debug.log(). The logging system supports hierarchical per-channel flags. You can enable/disable '
        'logs for specific subsystems (e.g., disable "sensor.co" to silence CO debug, disable "actuator.*" to silence all actuator logs). Remote UDP logging '
        'is optional; logs can be sent to a listener on the network for centralized monitoring.'
    )

    # 4. Board B (Actuator Board) – Complete Module Guide
    doc.add_heading('4. Board B (Actuator Board) – Complete Module Guide', level=1)
    doc.add_paragraph(
        'Board B is the "hands" of the system. It receives state from Board A and drives physical actuators to provide feedback, visualization, and control. '
        'The architecture mirrors Board A (main loop, timers, state, non-blocking update pattern) but the modules are actuator drivers instead of sensor drivers.'
    )

    doc.add_heading('4.1 Main Entry Point and Initialization', level=2)
    doc.add_paragraph(
        'Like Board A, main.py orchestrates boot: OTA check, WiFi, logging, buttons/actuators init (or simulation), ESP-NOW (now as server), UDP listener, '
        'then main loop.'
    )
    doc.add_paragraph(
        'Board B is the ESP-NOW server, meaning it listens for incoming messages from Board A. Once a link is established, Board B can receive sensor data '
        'and send back actuator state.'
    )

    doc.add_heading('4.2 Actuator Loop (core.actuator_loop)', level=2)
    doc.add_paragraph(
        'Similar to sensor_loop on Board A, actuator_loop is called every iteration and coordinates all actuator updates. It calls driver update functions '
        'with internal timers: LED blinking update (every 50 ms?), servo smoothing (every 50 ms), LCD refresh, buzzer pattern progression, etc.'
    )
    doc.add_paragraph(
        'Each actuator driver manages its own runtime state (e.g., servo current angle, LED blink phase, buzzer tone phase). The loop just tells each driver '
        'to update itself; the driver figures out what to do based on its internal timers and state.'
    )

    doc.add_heading('4.3 LED Control (actuators.leds)', level=2)
    doc.add_paragraph(
        'Three LEDs (green, blue, red) are digital outputs on GPIO16, GPIO17, GPIO19. Each LED can be off, on, or blinking. The blinking mode uses a cycle time '
        '(total period) and duty (time on vs. off) per cycle. A runtime state machine tracks which phase each LED is in and updates the GPIO accordingly on '
        'every update() call.'
    )
    doc.add_paragraph(
        'On power-up, green is always on (system ready indicator), blue and red are off. When an alarm is triggered, red LED is set to steady on (danger) or '
        'blinking (warning). When alarm clears, red goes off again. This visual feedback is crucial for users to know the system status at a glance.'
    )

    doc.add_heading('4.4 Servo Gate Control (actuators.servo)', level=2)
    doc.add_paragraph(
        'The servo (SG90) is a 9-gram servo motor that controls a gate. It is driven by a PWM signal on GPIO23 at 50 Hz frequency. The pulse width determines '
        'the angle: 0.5 ms = 0°, 2.5 ms = 180°. The driver uses a PWM peripheral to generate this signal.'
    )
    doc.add_paragraph(
        'The servo can move in two ways: smooth (gradual ramp to target angle over multiple loop cycles) or immediate (jump to target angle instantly). The smooth '
        'mode is used for the gate opening/closing animation; immediate mode can be used for emergency situations.'
    )
    doc.add_paragraph(
        'Gate automation: when the ultrasonic sensor on Board A detects presence (<50 cm), it sends this flag to Board B. Board B then opens the servo to 90° '
        '(gate open). When presence is lost, Board B starts a timer (delay before closing, e.g., 10 seconds). After the delay, the servo closes (0°). This '
        'automation runs independently in the update() function and respects both the presence state and timing.'
    )

    doc.add_heading('4.5 LCD Display (actuators.lcd)', level=2)
    doc.add_paragraph(
        'The 1602 LCD (16 characters × 2 lines) displays system status. It is connected via I2C (GPIO21/22) to a standard I2C backpack. The driver provides '
        'functions to display custom text or restore default status lines. Non-blocking operations: any write to the LCD is fast enough (a few milliseconds) '
        'that it does not need to be split across multiple loop iterations.'
    )
    doc.add_paragraph(
        'By default, the LCD shows "System Ready" and "Standby...". When an alarm is triggered, it can show alarm details (e.g., "CO HIGH" and "DANGER"). '
        'Users can also send commands to display custom messages.'
    )

    doc.add_heading('4.6 Buzzer (actuators.buzzer)', level=2)
    doc.add_paragraph(
        'The passive buzzer (DFRobot) is driven by PWM on GPIO25. Different frequencies produce different pitches; different duty cycles and patterns produce '
        'different sounds. The driver defines sound patterns (warning beep pattern, danger alarm pattern) as lists of (duration, is_tone, frequency) tuples.'
    )
    doc.add_paragraph(
        'On each update() call, the driver advances the current sound pattern based on elapsed time. When the system is in warning state, the warning pattern plays '
        'continuously. In danger state, the danger pattern (longer, more urgent beeps) plays. In normal state, buzzer is silent. Patterns loop; the update() function '
        'does not block, just progresses the phase counter.'
    )

    doc.add_heading('4.7 Audio (DFPlayer Mini)', level=2)
    doc.add_paragraph(
        'The DFPlayer Mini is an audio playback module connected via UART1 (GPIO27 TX, GPIO26 RX). It plays audio files stored on a microSD card. Commands are '
        'sent as serial frames. The driver provides play(), stop(), and set_volume() functions. This can be used for voice prompts ("Danger: high CO") or musical '
        'alarms. Implementation is simple: send a command frame, no need to wait for response; the module plays asynchronously.'
    )

    doc.add_heading('4.8 Buttons and SOS (actuators.buttons)', level=2)
    doc.add_paragraph(
        'Board B has a physical button (GPIO18) that users can press for emergency SOS. When pressed, an emergency mode is triggered: red LED flashes rapidly, '
        'buzzer sounds continuously, and the event is sent to Board A via ESP-NOW for logging and cloud notification. This is a critical safety feature.'
    )

    doc.add_heading('4.9 State (core.state)', level=2)
    doc.add_paragraph(
        'Board B also has core/state.py, but with different fields: actuator_state (LED modes, servo angle, buzzer/audio status), received_sensor_state (copy '
        'of sensor data from Board A + stale flags), system_control, and button_state.'
    )

    doc.add_heading('4.10 Communication: ESP-NOW Server', level=2)
    doc.add_paragraph(
        'Board B is the ESP-NOW server. It listens for incoming messages from Board A, updates its received_sensor_state, sends back an ACK and current '
        'actuator state, and continues normal operation. The ESP-NOW update() is non-blocking; reception happens asynchronously in interrupt handlers.'
    )

    # 5. Hardware Pins and Wiring
    doc.add_heading('5. Hardware Pins, Wiring, and Pin Selection Rationale', level=1)
    doc.add_paragraph(
        'ESP32 has many GPIO pins, but not all are equally suitable. Some pins are "strapping pins" (used at boot to determine boot mode), some conflict '
        'with WiFi or JTAG, and some have no ADC or PWM support. Let us review the chosen pins and why.'
    )

    doc.add_heading('5.1 Board A (Sensor Board) Pins', level=2)
    _add_table(doc, [
        ('DS18B20 temperature', 'GPIO4', 'Single-wire compatible; not a strap pin; I2C bus (21/22) free for other use.'),
        ('MQ-7 CO analog', 'GPIO35 (ADC1_7)', 'ADC1 works during WiFi (unlike ADC2); isolated from WiFi channels; stable supply.'),
        ('Buttons b1/b2/b3', 'GPIO16 / 17 / 19', 'Digital inputs, no conflict with WiFi; not strap pins; grouped for clarity.'),
        ('Ultrasonic TRIG', 'GPIO5', 'Output-capable, no conflict.'),
        ('Ultrasonic ECHO', 'GPIO18', 'Input-capable, no conflict.'),
        ('MAX30102 I2C SDA', 'GPIO21', 'Standard ESP32 I2C pin; 3.3V safe; shared bus OK.'),
        ('MAX30102 I2C SCL', 'GPIO22', 'Standard ESP32 I2C pin; 3.3V safe.'),
    ])
    doc.add_paragraph(
        'Why GPIO4 for DS18B20 and not GPIO2 or GPIO12? GPIO2 and GPIO12 are strap pins; using them can cause boot issues. GPIO4 is reliable and '
        'isolated. Why GPIO35 for CO? ADC2 conflicts with WiFi; ADC1 is safe. GPIO35 is ADC1_7, plenty of range.'
    )

    doc.add_heading('5.2 Board B (Actuator Board) Pins', level=2)
    _add_table(doc, [
        ('LED Green', 'GPIO16', 'Digital output; no conflict.'),
        ('LED Blue', 'GPIO17', 'Digital output; no conflict.'),
        ('LED Red', 'GPIO19', 'Digital output; no conflict.'),
        ('Servo PWM', 'GPIO23', 'PWM-capable; 50 Hz support; no conflict.'),
        ('LCD I2C SDA', 'GPIO21', 'Standard I2C; shared with Board A for consistency.'),
        ('LCD I2C SCL', 'GPIO22', 'Standard I2C.'),
        ('Buzzer PWM', 'GPIO25', 'PWM-capable; isolated from other use.'),
        ('DFPlayer TX', 'GPIO27', 'UART1 TX; 3.3V logic.'),
        ('DFPlayer RX', 'GPIO26', 'UART1 RX; 3.3V logic.'),
        ('Button/SOS', 'GPIO18', 'Input; not strap; interrupt-friendly.'),
    ])
    doc.add_paragraph(
        'Why the same GPIO numbering for LEDs and I2C? For consistency and mental ease. The high number for the buzzer and DFPlayer allows future expansion '
        'on lower numbers without changing existing assignments. All power pins (5V, 3.3V, GND) must be adequately sized for all loads (servos and audio '
        'can draw 500 mA+ peaks).'
    )

    # 6. Boot Sequence
    doc.add_heading('6. Initialization Flow and Boot Sequence', level=1)
    doc.add_paragraph(
        'Understanding the boot sequence is critical for debugging startup issues. Here is the complete sequence on both boards.'
    )
    doc.add_heading('6.1 Phase 0: Power-On and Hardware Reset', level=2)
    doc.add_paragraph(
        'When the board is powered on, the ESP32 hardware runs a bootloader, loads firmware from flash, and begins executing main.py. This all happens '
        'before Python code runs.'
    )
    doc.add_heading('6.2 Phase 1: OTA Update Check', level=2)
    doc.add_paragraph(
        'The very first line of main.py imports and calls ota_update.check_and_update(). This function checks if a new firmware version is available and '
        'downloads/installs it if needed. If an update is done, the board reboots and the new firmware starts. This ensures the latest code is always running '
        'on first power-on or after an update request.'
    )
    doc.add_heading('6.3 Phase 2: Module Imports and Debug Setup', level=2)
    doc.add_paragraph(
        'main.py imports all modules: debug, core, communication, config, etc. Debug logging is set up with initial flags (e.g., disable verbose logs, enable '
        'only critical channels). This allows debugging boot messages without spam.'
    )
    doc.add_heading('6.4 Phase 3: WiFi Connection', level=2)
    doc.add_paragraph(
        'core.wifi.init_wifi() is called. It reads WiFi credentials from config, connects to the network, and waits (with timeout) until connected or timeout '
        'occurs. This is a blocking operation, but it times out quickly if the network is unreachable. A log line reports success or failure.'
    )
    doc.add_heading('6.5 Phase 4: Remote Logging Initialization', level=2)
    doc.add_paragraph(
        'debug.init_remote_logging() sets up UDP logging sink if enabled. If disabled or fails, it is silently skipped. This allows all subsequent logs to go '
        'to the network if desired.'
    )
    doc.add_heading('6.6 Phase 5: Hardware Initialization or Simulation Mode', level=2)
    doc.add_paragraph(
        'If simulate_sensors (Board A) or simulate_actuators (Board B) is true, the simulation module is loaded. Otherwise, each sensor/actuator driver is '
        'initialized. If a driver fails to initialize (e.g., sensor not found on I2C), a warning is logged and operation continues. This "fail soft" approach '
        'allows the system to work with partial hardware available.'
    )
    doc.add_heading('6.7 Phase 6: ESP-NOW Communication Setup', level=2)
    doc.add_paragraph(
        'Board A (client) and Board B (server) initialize ESP-NOW. Board A begins trying to register Board B as a peer; Board B listens. A link is established '
        'once both are running. If one board boots first, it will retry periodically until the other comes online.'
    )
    doc.add_heading('6.8 Phase 7: UDP Command Listener', level=2)
    doc.add_paragraph(
        'A UDP socket is opened in non-blocking mode on port 37022. Test scripts can immediately start sending commands.'
    )
    doc.add_heading('6.9 Phase 8 (Board A only): Node-RED / MQTT Bridge', level=2)
    doc.add_paragraph(
        'Board A attempts to connect to the MQTT broker. If the broker is unreachable, a warning is logged and a reconnect timer is started. The system does '
        'not wait; it continues. MQTT will be retried every few seconds in the main loop.'
    )
    doc.add_heading('6.10 Phase 9: Main Loop Starts', level=2)
    doc.add_paragraph(
        'All initialization is complete. The main loop begins running at full CPU speed, calling update() functions in sequence. From this point on, timers '
        'control all periodic work. No more blocking waits.'
    )

    # 7. Main Loop Architecture
    doc.add_heading('7. Main Loop Architecture and Non-Blocking Design', level=1)
    doc.add_paragraph(
        'The main loop is the heartbeat of the system. Understanding its structure and timing is key to working with the firmware.'
    )
    doc.add_heading('7.1 Loop Structure', level=2)
    doc.add_paragraph(
        'while True: (infinite loop)'
    )
    doc.add_paragraph('    Check system control flags (e.g., reboot request)', style='List Bullet')
    doc.add_paragraph('    Call sensor_loop.update() / actuator_loop.update()', style='List Bullet')
    doc.add_paragraph('    Call espnow_communication.update()', style='List Bullet')
    doc.add_paragraph('    Call nodered_client.update() (Board A only)', style='List Bullet')
    doc.add_paragraph('    Call udp_commands.update()', style='List Bullet')
    doc.add_paragraph('    (No sleep; loop restarts immediately)', style='List Bullet')

    doc.add_heading('7.2 Why No Sleep?', level=2)
    doc.add_paragraph(
        'Many IoT firmware use sleep(100) or similar to "pace" the main loop and save CPU. But that blocks the CPU for 100 ms, delaying command reception, '
        'communication, and sensor responsiveness. Our design uses elapsed() timers instead: each function returns immediately if its work is not due, so the '
        'loop spins at full speed. This means messages and commands are handled within milliseconds, not 100 ms delays.'
    )
    doc.add_paragraph(
        'The tradeoff is CPU power. The ESP32 will run hot if the loop spins forever. To reduce power consumption, the loop could call a very brief yield() or '
        'feed the watchdog, but no long blocking sleep. (In a future optimization, we could use interrupt-driven designs to wake on events.)'
    )

    doc.add_heading('7.3 Elapsed() Timer Pattern', level=2)
    doc.add_paragraph(
        'The elapsed() function is called at the start of each periodic task to decide whether to proceed. Example: if elapsed("co", 1000): read_co(). '
        'The first time this is called, elapsed() returns True and records a timestamp. Until 1000 ms later, elapsed() returns False. After 1000 ms, it '
        'returns True again and updates the timestamp. This pattern repeats infinitely. Each task has its own key (string identifier), so timers do not '
        'interfere with each other.'
    )
    doc.add_paragraph(
        'This is extremely simple but effective. It means the code is easy to understand: if_elapsed( "X", interval ): do_X(). No complex scheduling needed.'
    )

    doc.add_heading('7.4 Loop Iteration Times', level=2)
    doc.add_paragraph(
        'How long does one iteration of the loop take? Typically 1–10 milliseconds, depending on how much work is being done. Most of the time, work is NOT due '
        '(elapsed() returns False immediately), so iterations are very short (< 1 ms). Occasionally, a sensor read happens (adds a few ms), or a communication '
        'update happens (adds a few ms). On average, loop iterations are extremely fast, leaving the system responsive.'
    )

    # 8. Communication Protocols
    doc.add_heading('8. Communication Protocols and Data Flow', level=1)
    doc.add_paragraph(
        'Three communication channels carry data through the system: ESP-NOW between boards, MQTT to the cloud, and UDP for local testing.'
    )

    doc.add_heading('8.1 ESP-NOW (Board A to Board B and Back)', level=2)
    doc.add_paragraph(
        'ESP-NOW is Espressif\'s proprietary wireless protocol for direct board-to-board communication without WiFi. It is faster and lower latency than WiFi. '
        'In our system, it carries sensor readings from A to B and actuator state from B to A.'
    )
    doc.add_paragraph(
        'Flow:'
    )
    doc.add_paragraph('Board A: Every 2.5 seconds, package (temp, CO, HR, SpO2, distance, presence, alarm level) and send to B.', style='List Number')
    doc.add_paragraph('Board B: Receive message, update state, send ACK + current actuator state back to A.', style='List Number')
    doc.add_paragraph('Board A: Receive ACK + actuator state, update received_actuator_state.', style='List Number')
    doc.add_paragraph('Both: If no message/ACK for >15 seconds, mark other board as stale (link broken).', style='List Number')
    doc.add_paragraph(
        'Message format is binary (not JSON) for efficiency. Typical payload ~50 bytes. Latency is ~10–50 ms. Reliability is good; ESP-NOW has built-in '
        'retry and ACK. If a message fails after retries, it is discarded (not guaranteed delivery, but good enough for real-time feedback).'
    )

    doc.add_heading('8.2 MQTT to Node-RED and App (Board A Only)', level=2)
    doc.add_paragraph(
        'Board A is the gateway to the cloud. It publishes system state to an MQTT broker (e.g., Adafruit IO) on a feed (topic) called "esp-a-to-nodered". '
        'Node-RED subscribes to this feed and can display the data on dashboards, log it to a database, or send alerts. A mobile app can also subscribe to '
        'the same feed to show live status to users.'
    )
    doc.add_paragraph(
        'Node-RED can also send commands back to Board A on a feed called "nodered-to-esp-a". Board A receives these commands and either executes them locally '
        'or forwards them to Board B via ESP-NOW. This allows remote control of actuators from the app or Node-RED dashboard.'
    )
    doc.add_paragraph(
        'MQTT uses TCP, so reliability is guaranteed: messages do not get lost (if a message is published while the broker is down, the broker buffers it). '
        'Latency is 100–500 ms depending on internet quality. This is acceptable for dashboard/app updates but too slow for real-time control (that is why '
        'ESP-NOW is used locally).'
    )
    doc.add_paragraph(
        'The MQTT connection is optional: if the broker is down or unreachable, Board A logs a warning and retries every 5 seconds. The rest of the system '
        '(local sensing, esp-now, actuation) continues unaffected.'
    )

    doc.add_heading('8.3 UDP Commands (Both Boards)', level=2)
    doc.add_paragraph(
        'For testing and debugging, both boards listen for UDP commands on port 37022. A Python script (send_command.py) sends JSON commands over the local network. '
        'Example: {"target":"A", "command":"simulate", "args":["co", "150"]} to fake a CO spike. UDP is fast and simple (no connection setup), but not reliable '
        '(packets can be lost). It is ideal for testing; not recommended for production control.'
    )

    # 9. Alarm Logic Deep Dive
    doc.add_heading('9. Alarm Logic System – Detailed Walkthrough', level=1)
    doc.add_paragraph(
        'The alarm logic is the "brain" that turns raw sensor readings into actionable states. Let us trace through a complete example to understand how it works.'
    )

    doc.add_heading('9.1 Sensor Thresholds and Time Windows', level=2)
    doc.add_paragraph(
        'Every monitored sensor has configurable thresholds and time windows (in config.json):'
    )
    doc.add_paragraph('CO: critical threshold = 50 ppm; warning window = 5s; danger window = 30s; recovery window = 10s', style='List Bullet')
    doc.add_paragraph('Temperature: safe range = 10–35°C; warning window = 10s; danger window = 60s; recovery window = 15s', style='List Bullet')
    doc.add_paragraph('Heart Rate: safe range = 50–120 BPM; warning window = 10s; danger window = 60s; recovery window = 15s', style='List Bullet')
    doc.add_paragraph('SpO2: safe threshold = 90%; warning window = 10s; danger window = 60s; recovery window = 15s', style='List Bullet')
    doc.add_paragraph(
        'These are all configurable, so you can adjust sensitivity without recompiling code. Shorter windows = faster reaction time but more false alarms. '
        'Longer windows = slower reaction but fewer false positives.'
    )

    doc.add_heading('9.2 State Machine Walkthrough: CO Sensor Example', level=2)
    doc.add_paragraph(
        'Let us trace a CO alarm from start to finish. Assume default config: critical = 50 ppm, warning window = 5s, danger window = 30s, recovery window = 10s.'
    )
    doc.add_paragraph('t=0s: System power-on. CO sensor reads 20 ppm (normal). Alarm state for CO = "normal", timer = NULL.', style='List Number')
    doc.add_paragraph('t=5s: CO sensor reads 60 ppm (above 50 ppm threshold). Critical condition detected! Timer starts. Alarm state remains "normal" (waiting for warning timeout). State.system_state["co_level"] = "normal" (not updated yet).', style='List Number')
    doc.add_paragraph('t=8s: CO sensor still reads 60 ppm. Timer at 3s (< 5s warning window). Alarm state still "normal".', style='List Number')
    doc.add_paragraph('t=10.5s: CO sensor reads 60 ppm. Timer at 5.5s (> 5s warning window). Alarm state changes to "warning". Buzzer starts warning pattern. Red LED starts blinking. state.system_state["co_level"] = "warning".', style='List Number')
    doc.add_paragraph('t=12s: CO sensor reads 65 ppm. Timer keeps running. Alarm state still "warning".', style='List Number')
    doc.add_paragraph('t=35s: CO sensor reads 62 ppm. Timer at 30s (>= 30s danger window). Alarm state changes to "danger". Buzzer switches to danger pattern (urgent). Red LED becomes solid on. state.system_state["co_level"] = "danger".', style='List Number')
    doc.add_paragraph('t=60s: CO sensor reads 45 ppm (below 50 ppm). Critical condition cleared. Timer stops. Recovery window starts (10s). Alarm state remains "danger" (not immediately downgraded).', style='List Number')
    doc.add_paragraph('t=70s: Recovery timer elapsed (10s). Alarm state returns to "normal". Buzzer stops. Red LED turns off. state.system_state["co_level"] = "normal".', style='List Number')
    doc.add_paragraph(
        'This state machine prevents fluttering (rapid on/off switching) by using time windows. A single glitch in a sensor reading does not immediately trigger '
        'an alarm; the condition must persist. Recovery also uses a window to avoid immediate downgrade when a condition briefly clears.'
    )

    doc.add_heading('9.3 Global Alarm Level', level=2)
    doc.add_paragraph(
        'Each sensor (CO, temp, HR, SpO2) has its own alarm level (normal/warning/danger). The global alarm level is the worst of all:'
    )
    doc.add_paragraph('If any sensor is "danger" → global alarm = "danger"', style='List Bullet')
    doc.add_paragraph('Else if any sensor is "warning" → global alarm = "warning"', style='List Bullet')
    doc.add_paragraph('Else → global alarm = "normal"', style='List Bullet')
    doc.add_paragraph(
        'This is updated on every alarm_logic.update() call and stored in state.alarm_state["level"]. Board B sees this global level and acts accordingly: red LED '
        'solid for danger, blinking for warning, off for normal. If multiple sensors are in danger, the system is clearly in danger, even if one recovers.'
    )

    doc.add_heading('9.4 Actuator Reactions', level=2)
    doc.add_paragraph(
        'When the alarm state changes, Board B reacts visually and aurally:'
    )
    doc.add_paragraph('Danger: Red LED solid on. Buzzer plays urgent danger pattern continuously. LCD shows alarm details.', style='List Bullet')
    doc.add_paragraph('Warning: Red LED blinks slowly. Buzzer plays warning pattern (repeating beeps). LCD shows alert.', style='List Bullet')
    doc.add_paragraph('Normal: Red LED off. Buzzer silent. LCD shows default text (System Ready).', style='List Bullet')
    doc.add_paragraph(
        'Importantly, these reactions are driven by the alarm state received from Board A via ESP-NOW every 2.5s. So there is a small delay (up to 2.5s) between '
        'when a sensor reading goes critical on Board A and when the actuators react on Board B. But this is acceptable for health monitoring; it is not real-time '
        'safety-critical control.'
    )

    # 10. State Management
    doc.add_heading('10. State Management and Shared Data Structures', level=1)
    doc.add_paragraph(
        'State is centralized in core/state.py on each board. Every module reads and writes from this single source of truth. Let us see the structure.'
    )

    doc.add_heading('10.1 Board A State Structure', level=2)
    doc.add_paragraph(
        'sensor_data = { "temperature": float, "co": float, "heart_rate": { "ir": int, "red": int, "bpm": int, "spo2": int, "status": str }, '
        '"ultrasonic_distance_cm": float, "ultrasonic_presence": bool, "acc": { "x": float, "y": float, "z": float } }'
    )
    doc.add_paragraph(
        'system_state = { "co_level": "normal/warning/danger", "temp_level": "...", "heart_level": "..." }'
    )
    doc.add_paragraph(
        'alarm_state = { "level": "normal/warning/danger", "source": str (which sensor?) }'
    )
    doc.add_paragraph(
        'button_state = { "b1": bool, "b2": bool, "b3": bool }'
    )
    doc.add_paragraph(
        'received_actuator_state = { "leds": { "green": bool, "blue": bool, "red": bool }, "servo": { "angle": int, "moving": bool }, ... }'
    )
    doc.add_paragraph(
        'system_control = { "reboot_requested": bool, ... }'
    )

    doc.add_heading('10.2 Board B State Structure', level=2)
    doc.add_paragraph(
        'actuator_state = { "leds": { "green": bool, ..., "led_modes": { "green": "off/on/blinking", ... } }, "servo": { "angle": int, "moving": bool }, '
        '"lcd": { "line1": str, "line2": str }, ... }'
    )
    doc.add_paragraph(
        'received_sensor_state = { "temperature": float, ..., "presence_detected": bool, "alarm_level": str, "last_update": timestamp, "is_stale": bool }'
    )
    doc.add_paragraph(
        'All modules on each board import state and read/write these dictionaries. There is no need to pass data between modules or worry about sync bugs; '
        'the central state is the single source of truth.'
    )

    # 11. Debug and Logging
    doc.add_heading('11. Debug, Logging, and Remote Monitoring', level=1)
    doc.add_paragraph(
        'Debugging a distributed system with two wireless boards is challenging. The debug system provides tools to observe and control the system.'
    )

    doc.add_heading('11.1 Hierarchical Log Flags', level=2)
    doc.add_paragraph(
        'All log messages go through debug.debug.log(channel, message). The channel name uses dots to create a hierarchy: "sensor.co.read", "actuator.led.set", '
        'etc. You can enable/disable logs by prefix: disable "sensor.*" to silence all sensors, disable "sensor.co" to silence CO only, etc.'
    )
    doc.add_paragraph(
        'This allows precise control: during normal operation, enable only critical logs (errors, state changes). During debugging, enable everything and filter '
        'to the subsystem you are interested in.'
    )

    doc.add_heading('11.2 Remote UDP Logging', level=2)
    doc.add_paragraph(
        'Logs can be sent to a UDP listener on the network for centralized monitoring. The board sends each log message over UDP; a Python listener script '
        '(log_listener.py) receives and prints them. This is useful when developing or troubleshooting: you can watch both boards\' logs in real time from your PC.'
    )
    doc.add_paragraph(
        'Remote logging is non-blocking: if the listener is not present or the network is down, logs are silently dropped. The system continues normally.'
    )

    doc.add_heading('11.3 Testing with UDP Commands', level=2)
    doc.add_paragraph(
        'For testing, send_command.py allows injecting commands over UDP. Examples:'
    )
    doc.add_paragraph('{"target":"A","command":"simulate","args":["co","100"]} – fake CO spike', style='List Bullet')
    doc.add_paragraph('{"target":"B","command":"led","args":["red","on"]} – turn red LED on', style='List Bullet')
    doc.add_paragraph('{"target":"B","command":"servo","args":["90"]} – move servo to 90°', style='List Bullet')
    doc.add_paragraph('{"target":"A","command":"status"} – request current state', style='List Bullet')
    doc.add_paragraph(
        'These commands are received by udp_commands.update(), parsed, and routed via command_handler to local handlers or ESP-NOW to the other board. '
        'This allows testing individual components without waiting for real sensor input.'
    )

    # 12. Simulation Mode
    doc.add_heading('12. Simulation Mode and Testing', level=1)
    doc.add_paragraph(
        'To develop and test without physical hardware, enable simulation mode in config.json: "simulate_sensors": true (Board A) or "simulate_actuators": true (Board B). '
        'The firmware then loads simulation modules that generate synthetic data instead of reading real hardware.'
    )
    doc.add_paragraph(
        'For Board A, simulated sensors produce random values in realistic ranges: temperature 20–30°C, CO 0–50 ppm, HR 60–100 BPM, distance 50–200 cm. '
        'The alarm logic works normally, triggering on spikes. Actuators on Board B react as normal.'
    )
    doc.add_paragraph(
        'For Board B, simulated actuators accept commands but do not drive GPIO. LEDs, servo, buzzer, LCD all pretend to work: state is updated, but no physical '
        'output happens. This is perfect for testing logic and communication without hardware damage.'
    )

    # 13. Node-RED Integration
    doc.add_heading('13. Node-RED Integration and MQTT Topics', level=1)
    doc.add_paragraph(
        'Node-RED is a flow-based visual programming environment for IoT. Board A acts as the MQTT client, publishing state and consuming commands. Here is how '
        'to set up Node-RED to interact with the system.'
    )

    doc.add_heading('13.1 MQTT Topics and Message Format', level=2)
    doc.add_paragraph(
        'Board A publishes to: "<username>/feeds/esp-a-to-nodered" (state/telemetry). Message is a JSON string with all sensor readings and alarm level.'
    )
    doc.add_paragraph(
        'Board A subscribes to: "<username>/feeds/nodered-to-esp-a" (commands). Message is a JSON string with target, command, and arguments.'
    )
    doc.add_paragraph(
        'Example state message:'
    )
    doc.add_paragraph(
        '{"temperature": 25.5, "co": 12.3, "heart_rate": {"bpm": 75, "spo2": 98}, "distance": 120, "presence": true, "alarm_level": "normal"}'
    )
    doc.add_paragraph(
        'Example command message:'
    )
    doc.add_paragraph(
        '{"target":"B", "command":"servo", "args":["90"]}'
    )

    doc.add_heading('13.2 Node-RED Dashboard Setup', level=2)
    doc.add_paragraph(
        'In Node-RED, create MQTT input nodes to subscribe to esp-a-to-nodered and parse the JSON. Display readings on gauges, graphs, and text nodes. '
        'Create MQTT output nodes to send commands when users click buttons or sliders on the dashboard. This gives non-technical users a nice web interface '
        'to view and control the system.'
    )

    # 14. Operational Procedures
    doc.add_heading('14. Operational Procedures and Troubleshooting', level=1)
    doc.add_heading('14.1 First-Time Startup', level=2)
    doc.add_paragraph(
        'Step 1: Power on Board B first (it is the ESP-NOW server). Wait for boot logs to appear (if serial monitor is connected).'
    )
    doc.add_paragraph(
        'Step 2: Power on Board A. It will attempt to establish ESP-NOW link with B. Watch logs for "ESP-NOW link established" or similar.'
    )
    doc.add_paragraph(
        'Step 3: Verify LED status: green should be on (Board B); red and blue off (unless an alarm is active). LCD should show default text.'
    )
    doc.add_paragraph(
        'Step 4: Use a network tool (ping, nslookup) to verify connectivity to the MQTT broker (if using Node-RED). Board A should attempt to connect every ~5s if not already connected.'
    )

    doc.add_heading('14.2 Testing Sensors', level=2)
    doc.add_paragraph(
        'Use send_command.py to inject fake sensor data and verify alarm logic. Example:'
    )
    doc.add_paragraph('python send_command.py A simulate co 100 – Injects CO = 100 ppm. Red LED should start blinking (warning) after 5s.', style='List Bullet')
    doc.add_paragraph('Watch logs to confirm state transitions. Check MQTT feed on Node-RED dashboard to see the spike transmitted to cloud.', style='List Bullet')

    doc.add_heading('14.3 Testing Actuators', level=2)
    doc.add_paragraph(
        'Use send_command.py to drive actuators. Example:'
    )
    doc.add_paragraph('python send_command.py B led red on – Turns red LED on.', style='List Bullet')
    doc.add_paragraph('python send_command.py B servo 90 – Moves servo to 90° (gate opens).', style='List Bullet')
    doc.add_paragraph('python send_command.py B lcd "Line1" "Line2" – Writes custom text to LCD.', style='List Bullet')

    doc.add_heading('14.4 Troubleshooting: ESP-NOW Link Down', level=2)
    doc.add_paragraph(
        'Symptom: Logs show "WARNING: Sensor data from B is stale" (on A) or vice versa after >15s.'
    )
    doc.add_paragraph(
        'Cause: ESP-NOW communication is broken. Possible reasons: boards out of range (unlikely indoors), WiFi on nearby channel interfering, hardware fault.'
    )
    doc.add_paragraph(
        'Solution: Power-cycle both boards. Check WiFi channel in config. If persistent, check antenna connections. Use a signal strength scanner to detect '
        'interference.'
    )

    doc.add_heading('14.5 Troubleshooting: MQTT Connection Failing', level=2)
    doc.add_paragraph(
        'Symptom: Logs show "MQTT connect failed" repeatedly.'
    )
    doc.add_paragraph(
        'Cause: Broker unreachable, credentials wrong, or firewall blocking.'
    )
    doc.add_paragraph(
        'Solution: Verify broker address and port in config.json. Verify WiFi credentials. Test connectivity from PC: '
        'mosquitto_pub -h <broker> -u <user> -P <pass> -t test -m "hello". If that fails, broker is unreachable from your network.'
    )

    doc.add_heading('14.6 Troubleshooting: Sensor Reads Stuck at NULL', level=2)
    doc.add_paragraph(
        'Symptom: state.sensor_data["temperature"] is always None.'
    )
    doc.add_paragraph(
        'Cause: Sensor initialization failed (not detected on I2C or OneWire) or driver has a bug.'
    )
    doc.add_paragraph(
        'Solution: Check boot logs for "Temperature initialization failed" or similar. Verify wiring: correct pin, power, ground. Try running sensor driver '
        'test script in isolation. Check if another device is using the same I2C address or GPIO pin.'
    )

    # 15. Code Examples
    doc.add_heading('15. Code Examples and Usage Patterns', level=1)
    doc.add_heading('15.1 Adding a New Sensor', level=2)
    doc.add_paragraph(
        'To add a new sensor (e.g., humidity sensor), follow this pattern:'
    )
    doc.add_paragraph('Create sensors/humidity.py with init_humidity() and read_humidity() functions.', style='List Number')
    doc.add_paragraph('In init_humidity(), initialize the hardware and return True/False.', style='List Number')
    doc.add_paragraph('In read_humidity(), read the sensor and update state.sensor_data["humidity"].', style='List Number')
    doc.add_paragraph('In core.sensor_loop.update(), add: if elapsed("humidity", HUMIDITY_INTERVAL): read_humidity().', style='List Number')
    doc.add_paragraph('In core.state, add "humidity": None to sensor_data dict.', style='List Number')
    doc.add_paragraph('Add humidity threshold and alarm logic to logic.alarm_logic if needed.', style='List Number')
    doc.add_paragraph('Add config entries for pin, interval, and thresholds in config.json.', style='List Number')

    doc.add_heading('15.2 Adding a New Actuator', level=2)
    doc.add_paragraph(
        'To add a new actuator (e.g., relay for pump control), follow this pattern:'
    )
    doc.add_paragraph('Create actuators/relay.py with init_relay() and control_relay(on_off) functions.', style='List Number')
    doc.add_paragraph('In core.actuator_loop.update(), call actuator state updates if needed.', style='List Number')
    doc.add_paragraph('In communication.command_handler, add case for "relay" command to call control_relay().', style='List Number')
    doc.add_paragraph('Add config entry for pin in config.json.', style='List Number')

    doc.add_heading('15.3 Sending a Command Programmatically', level=2)
    doc.add_paragraph(
        'To send a command from Python (e.g., from a scheduling script), use send_command.py as a template or import and call espnow_communication directly:'
    )
    doc.add_paragraph(
        'import socket, json; sock = socket.socket(); sock.sendto(json.dumps({"target":"B","command":"led","args":["red","on"]}).encode(), ("192.168.1.100", 37022))'
    )

    doc.add_heading('15.4 Reading State in a Script', level=2)
    doc.add_paragraph(
        'To poll board state from a PC script, subscribe to the MQTT feed or periodically request state via UDP:'
    )
    doc.add_paragraph(
        'import socket, json; sock = socket.socket(); sock.sendto(json.dumps({"target":"A","command":"status"}).encode(), ("192.168.1.100", 37022)); '
        'resp, _ = sock.recvfrom(1024); print(json.loads(resp.decode()))'
    )

    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This system demonstrates several best practices in embedded IoT design: separation of concerns (two boards), non-blocking architecture (timers instead of sleep), '
        'centralized state, modular drivers, and layered communication (local ESP-NOW, cloud MQTT, debug UDP). Understanding these principles will help you maintain, '
        'extend, and debug the system effectively. Welcome to the team!'
    )

    # Save document
    output_path = os.path.expanduser('~/OneDrive/Desktop/Python Files/First Project/scripts/documentation')
    os.makedirs(output_path, exist_ok=True)
    doc_path = os.path.join(output_path, 'IoT_System_Complete_Guide.docx')
    if os.path.exists(doc_path):
        try:
            os.remove(doc_path)
        except Exception:
            pass
    doc.save(doc_path)
    print(f"Documentation created successfully: {doc_path}")
    return doc_path


if __name__ == '__main__':
    create_documentation()
